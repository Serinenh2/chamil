"""Génération des documents Word (.docx) — même contenu que les PDF (section 39.4)."""
from io import BytesIO

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

NAVY = RGBColor(0x0B, 0x25, 0x45)
BLUE = RGBColor(0x1B, 0x5C, 0xB4)
BLUE_HEX = "1B5CB4"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def _shade_cell(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _add_field(paragraph, field_code):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = field_code
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def _add_page_number_footer(section):
    footer = section.footer
    p = footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Page : ")
    run.font.size = Pt(7.5)
    _add_field(p, "PAGE")
    run2 = p.add_run(" / ")
    run2.font.size = Pt(7.5)
    _add_field(p, "NUMPAGES")

DOC_TITLES = {
    "quote": "Proforma",
    "sales_order": "Commande client",
    "delivery": "Bon de livraison",
    "customer_invoice": "Facture",
    "purchase_order": "Commande fournisseur",
    "goods_receipt": "Bon de réception",
}


def render_document_word(document, kind, context):
    """Retourne les octets du .docx pour un document donné."""
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    company = context["company"]
    party = context["party"]
    settings_obj = context["settings"]
    owner = context["owner"]
    currency = settings_obj.currency if settings_obj else "DA"

    if company and company.logo:
        try:
            doc.add_picture(company.logo.path, height=Cm(1.6))
        except (ValueError, OSError):
            pass

    brand = doc.add_paragraph()
    run = brand.add_run(company.name if company else "")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = NAVY

    capital = doc.add_paragraph()
    run = capital.add_run(
        f"Au capital de {company.capital} {currency}" if company else "")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    ref_table = doc.add_table(rows=0, cols=2)
    ref_table.alignment = WD_TABLE_ALIGNMENT.RIGHT

    def ref_row(label, value):
        row = ref_table.add_row()
        row.cells[0].text = label
        row.cells[1].text = "" if value is None else str(value)

    ref_row("Référence", document.number)
    ref_row("Date", document.date.strftime("%Y-%m-%d") if document.date else "")
    ref_row("Montant", f"{document.amount_ttc} {currency}")
    ref_row("Client", party.name if party else "")
    if context.get("reference_date"):
        ref_row(context["reference_date_label"], context["reference_date"].strftime("%Y-%m-%d"))

    title_table = doc.add_table(rows=1, cols=1)
    title_cell = title_table.rows[0].cells[0]
    _shade_cell(title_cell, BLUE_HEX)
    title_p = title_cell.paragraphs[0]
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(f"{DOC_TITLES.get(kind, 'Document')} N° : {document.number}")
    title_run.font.size = Pt(14)
    title_run.bold = True
    title_run.font.color.rgb = WHITE
    doc.add_paragraph()

    doc.add_paragraph(f"Code client : {getattr(party, 'code', '') or '—'}")
    doc.add_paragraph(f"N° RC : {getattr(party, 'rc', '') or '—'}")
    doc.add_paragraph(f"NIF : {getattr(party, 'nif', '') or '—'}")
    doc.add_paragraph(f"N° I.S. : {getattr(party, 'nis', '') or '—'}")
    doc.add_paragraph(f"Tél : {getattr(party, 'phone', '') or '—'}")
    dest = doc.add_paragraph()
    dest.add_run("Destinataire : ").bold = True
    dest.add_run(f"{party.name if party else ''} — {getattr(party, 'address', '') or ''}")

    headers = ["Réf. article", "Désignation", "Unité", "Quantité",
               "Prix U. HT", "Montant HT", "Remise", "TVA"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, label in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = label
        cell.paragraphs[0].runs[0].bold = True
    for line in context["lines"]:
        row = table.add_row()
        row.cells[0].text = line.product.code
        row.cells[1].text = line.description
        row.cells[2].text = line.product.unit or ""
        row.cells[3].text = str(line.quantity)
        row.cells[4].text = str(line.unit_price)
        row.cells[5].text = str(line.amount_ht)
        row.cells[6].text = f"{line.discount} %"
        row.cells[7].text = f"{line.vat_rate} %"

    doc.add_paragraph()
    totals = doc.add_table(rows=0, cols=2)

    def total_row(label, value, bold=False):
        row = totals.add_row()
        row.cells[0].text = label
        row.cells[1].text = f"{value} {currency}"
        if bold:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for r in para.runs:
                        r.bold = True

    total_row("Total HT", document.amount_ht)
    total_row("TVA", document.amount_vat)
    total_row("Total TTC", document.amount_ttc, bold=True)
    if kind == "customer_invoice":
        total_row("Montant payé", document.paid_amount)
        total_row("Net à payer", document.balance, bold=True)

    words = doc.add_paragraph()
    label = "la présente facture" if kind == "customer_invoice" else "le présent document"
    run = words.add_run(f"Arrêtée {label} à la somme de : {context['amount_words']}.")
    run.italic = True

    doc.add_paragraph()
    sig_table = doc.add_table(rows=1, cols=2)
    stamp_cell, sig_cell = sig_table.rows[0].cells

    if company and company.stamp:
        try:
            stamp_cell.paragraphs[0].add_run().add_picture(company.stamp.path, height=Cm(2))
        except (ValueError, OSError):
            pass
    stamp_p = stamp_cell.add_paragraph("Cachet de l'entreprise")
    stamp_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sig_cell_p1 = sig_cell.paragraphs[0]
    sig_cell_p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if owner and owner.signature:
        try:
            sig_cell_p1.add_run().add_picture(owner.signature.path, height=Cm(1.5))
        except (ValueError, OSError):
            pass
    name_p = sig_cell.add_paragraph(owner.full_name if owner else "")
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if owner:
        position_p = sig_cell.add_paragraph(owner.get_position_display())
        position_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.text = (
        f"RC {getattr(company, 'rc', '')} · NIF {getattr(company, 'nif', '')} · "
        f"N° I.S. {getattr(company, 'nis', '')} · Tél {getattr(company, 'phone', '')}"
    )
    _add_page_number_footer(section)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
